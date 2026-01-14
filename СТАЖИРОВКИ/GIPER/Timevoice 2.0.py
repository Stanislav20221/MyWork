import streamlit as st
from docx import Document
from dotenv import load_dotenv
import openai
from openai import OpenAI
import os
import zipfile
import rarfile
import shutil
import docx
import pandas as pd
import matplotlib.pyplot as plt
import seaborn as sns
from io import BytesIO
load_dotenv()

client = OpenAI(
    api_key=os.getenv("OPENAI_API_KEY")

)
themes = {'Консультации по товарам', 'Использование приобретённого изделия', 'Проблема, связанная с купленным товаром', 'Информация о заказе', 'Оформление заказа', 'Отмена заказа', 'Проблема, связанная с доставкой', 'Информация о доставке товара', 'Возврат/Обмен', 'Адрес АСЦ', 'Закупка комплектующих АСЦ', 'Нецелевой звонок',
           'Программа лояльности', 'Юридическое лицо', 'Жалоба', 'Рассылка/реклама'}
themecount = {}
results = []

dialog = '''
    1.  Консультации по товарам (клиент хочет приобрести товар, и спрашивает отличия/характеристики/стоимость и пр).
    2.  Использование приобретённого изделия (клиент уже купил товар, и его интересуют вопросы как использовать/включить/обслуживать приобретенный товар).
    3.  Проблема, связанная с купленным товаром (консультация по изделию, если возникла какая-либо проблема, связанная с ним – не включается, что-то не работает т т.п.).
    4.  Информация о заказе (клиент хочет уточнить уточить детали заказа, например, состав, способ доставки и т. п.).
    5.  Оформление заказа (оператор оформляет заказ в разговоре).
    6.  Отмена заказа (клиент хочет отменить заказ).
    7.  Проблема, связанная с доставкой (клиент жалуется на доставку, курьер не смог связаться и пр. Все, что требует от КЦ направить запрос в ТК или логистам.).
    8.  Информация о доставке товара (клиент интересуется датой, временем и сроками доставки товара, спрашивает, где и как забрать заказ и т. п.).
    9.  Возврат/Обмен (консультации по возврату и обмену, в т.ч. звонки, которые были переданы на Отдел Клиентского Сервиса ООО "Гипер").
    10.  Адрес АСЦ (если клиента интересует адрес АСЦ, или ему нужно сдать товар на ремонт).
    11.  Закупка комплектующих АСЦ (клиента интересует вопрос покупки комплектующих, если их нельзя купить в ИМ и купить можно в АСЦ).
    12.   Нецелевой звонок (звонки, которые не попадают в ЗО КЦ ООО "Гипер").
    13.  Программа лояльности (Вопросы по программе лояльности, списаниям баллов, возвратам баллов, сертификатам, бонусам и т. п.).
    14.  Юридическое лицо (если клиент является юридическим лицом – компанией, ООО, ИП, и т. п.).
    15.  Жалоба (клиент жалуется на действия сотрудников ООО "Гипер" или Транспортных Компаний).
    16.  Рассылка/реклама (речь идет о рассылке на почту/смс/etc.)
'''
UNRAR_PATH = r"C:\Program Files\WinRAR\UnRAR.exe"
if os.path.exists(UNRAR_PATH):
    rarfile.UNRAR_TOOL = UNRAR_PATH
else:
    st.error(f"Ошибка: UnRAR.exe не найден! Проверьте путь: {UNRAR_PATH}")
    st.stop()
st.set_page_config(page_title="Определение темы клиента", layout="wide")     
#==================================================================================== ФУНКЦИЯ ОТВЕТА CHAT GPT ======================================================================================================================================================
def answer_index(system, topic, dialog, temp, verbose=0):
    # Поиск релевантных отрезков из базы знаний по вопросу пользователя
    try:
        messages = [  # <-- Теперь тут 4 пробела
            {"role": "system", "content": system},
            {"role": "user", "content": f" Проанализируй текст диалога клиента м менеджера и выбери наиболее точную категорию из списка: {topic} Диалог {dialog}"}
        ]

        completion = client.chat.completions.create(  # <-- Уровень отступа теперь правильный
            model="gpt-4o-mini",
            messages=messages,
            temperature=temp
        )
        answer = completion.choices[0].message.content
        return answer
    except Exception as e:
        return "Ошибка доступа к Open AI. Проверьте API-ключ."
        

#==================================================================================== ФУНКЦИЯ РАСПАКОВКИ АРХИВА ======================================================================================================================================================
def unzip_file(file_bytes, filename, extract_to="unzipped_files"):
    """
    Распаковывает ZIP или RAR архив, переданный в виде байтового потока.

    :param file_bytes: Байтовый объект (BytesIO) с архивом.
    :param filename: Имя файла (для определения формата: ZIP или RAR).
    :param extract_to: Директория для извлечения файлов.
    :return: Путь к извлечённой папке или None при ошибке.
    """
    if not os.path.exists(extract_to):
        os.makedirs(extract_to)

    try:
        if filename.endswith(".zip"):
            with zipfile.ZipFile(file_bytes, 'r') as archive:
                archive.extractall(extract_to)
        elif filename.endswith(".rar"):
            with rarfile.RarFile(file_bytes) as archive:
                archive.extractall(extract_to)
        else:
            raise ValueError("Неподдерживаемый формат архива. Поддерживаются только ZIP и RAR.")

        print(f"Архив успешно распакован в: {extract_to}")
        return extract_to

    except (zipfile.BadZipFile, rarfile.BadRarFile):
        print("Ошибка: Файл повреждён или не является архивом.")
    except Exception as e:
        print(f"Произошла ошибка: {e}")

    return None

#==================================================================================== ФУНКЦИЯ ИЗВЛЕЧЕНИЯ ТЕКСТА ИЗ ФАЙЛОВ ======================================================================================================================================================

def extract_text_from_files(folder="unzipped_files"):
    """
    Извлекает текст из всех файлов в указанной папке.
    Поддерживает форматы: .txt, .docx, .pdf

    :param folder: Папка с распакованными файлами
    :return: Словарь {имя файла: извлеченный текст}
    """
    extracted_data = {}

    if not os.path.exists(folder):
        print("Ошибка: Папка с файлами не найдена!")
        return extracted_data

    for filename in os.listdir(folder):
        file_path = os.path.join(folder, filename)
        text = ""

        if filename.endswith(".docx"):
            doc = docx.Document(file_path)
            text = "\n".join([paragraph.text for paragraph in doc.paragraphs])

       # Добавляем текст в словарь
        if text:
            extracted_data[filename] = text

    return extracted_data



st.title("Определение темы обращения клиента")
tab1, tab2, tab3, tab4 = st.tabs(["Диалоги", "Результаты анализа", "Статистика (таблица)", "Статистика (график)"])
# Создание боковой панели
st.sidebar.header("Настройки модели")
temp = st.sidebar.slider("Температура", min_value=0.0, max_value=2.0, value=0.1) 
if "custom_css" not in st.session_state:
    st.session_state.custom_css = True  # Флаг, что CSS уже применён
    custom_text = "Перетащите файл сюда или выберите его вручную"
    st.markdown(f"""
        <style>
            .stFileUploader div div {{ visibility: hidden; }}
            .stFileUploader div div:after {{
                content: "{custom_text}";
                visibility: visible;
                display: block;
                text-align: center;
                font-size: 16px;
                color: #4F8BF9;
                font-weight: bold;
            }}
        </style>
    """, unsafe_allow_html=True)
#                                         ЗАГРУЗКА ФАЙЛА
st.sidebar.header("Загрузить файл")
uploaded_file = st.sidebar.file_uploader("Выберите файл для загрузки", type=['zip','rar'], key="file_uploader")
#if st.sidebar.button('Выбрать файл'):
if uploaded_file is not None:
    file_path = os.path.join("temp_archives", uploaded_file.name)

    # Создаём временную папку, если её нет
    if not os.path.exists("temp_archives"):
        os.makedirs("temp_archives")

    # Сохраняем файл на диск
    with open(file_path, "wb") as f:
        f.write(uploaded_file.getbuffer())

    # Распаковываем архив
    extracted_folder = unzip_file(file_path, uploaded_file.name)  

    if extracted_folder:  # Проверяем, что распаковка успешна
        st.sidebar.success(f"Файлы успешно извлечены в папку: {extracted_folder}")
        extracted_files = os.listdir(extracted_folder)  # Получаем список файлов
        st.sidebar.write("Содержимое архива:")
        st.sidebar.write(extracted_files)
        fcount = len(extracted_files)
        st.sidebar.write(f"Всего файлов, {len(extracted_files)}")
        shutil.rmtree("temp_archives")
    else:
        st.sidebar.error("Ошибка при распаковке архива!")


#     st.success("Диалог успешно загружен!")
#     if uploaded_file.name.endswith('.docx'):
#        doc = Document(uploaded_file)
#        content = "\n".join([para.text for para in doc.paragraphs])
#    #    st.text_area("Содержимое файла:", content, height=300)           
    
# #user_input = st.text_area("Введите ваше сообщение:", "")
    with tab1:
        st.write(extract_text_from_files())
    with tab2:
        if st.button("Анализировать диалог"):
            system = "Ты — виртуальный ассистент в магазине. Твоя задача — четко определить тему звонка покупателя на основе его обращения."
            et = extract_text_from_files()
            if (et):
               for index, (filename, text) in enumerate(et.items(), start=1):
                # st.write(f"Диалог {index}\n\n {text}")
                    topic = answer_index(system, text, dialog, temp)
                    if (topic == "Ошибка доступа к Open AI. Проверьте API-ключ."):
                        st.write(topic)
                        break
                    else:
                        results.append(topic)
                        st.write(f"Диалог {index}\n\n{topic}")
                        if index == len(et): 
                           st.write(f"Всего обработано файлов: {len(et)}")   
            for theme in themes:
                for result in results:
                    if theme.lower().strip() in result.lower().strip():  # Учитываем регистр и пробелы
                        themecount[theme] = themecount.get(theme, 0) + 1  # Правильное обновление словаря
        with tab3:
           df=pd.DataFrame(list(themecount.items()), columns=["Тема", "Количество"])
           df["%"] = (df["Количество"] / df["Количество"].sum() * 100).round(2)
           df_display = df.drop(columns=["%"])
           df.index = df.index + 1
           df_display.index.name = "№ п/п"
           st.table(df_display)
        with tab4:
        # Создаём фигуру и ось
            fig, ax = plt.subplots(figsize=(5, 2))

            # Строим график
            sns.barplot(x=df["Тема"], y=df["%"], hue=df["Тема"], palette="viridis", legend=False, ax=ax)

            # Настраиваем подписи
            for container in ax.containers:
                ax.bar_label(container, fmt="%.2f%%", fontsize=4, color="#FA8669", padding=3)       
            ax.set_xlabel("Тема обращения", fontsize=6, color="#FA8669")
            ax.set_ylabel("Процент обращений", fontsize=6, color="#FA8669")
            ax.set_title("Распределение тем обращений", fontsize=10, color="#FA8669")

            # Правильное задание подписей оси X
            ax.set_xticks(range(len(df["Тема"])))
            ax.set_xticklabels(df["Тема"], rotation=10, ha="right", fontsize=3, color="#FA8669")
            
            # Делаем фон оси тёмным
            ax.set_facecolor("#222222")  # Цвет фона внутри графика
            fig.patch.set_facecolor("#111111")  # Цвет фона всей фигуры

            # Делаем белыми оси и сетку
            ax.spines["bottom"].set_color("white")
            ax.spines["left"].set_color("white")
            ax.tick_params(axis="both", colors="white")


            # Отображаем график в Streamlit
            st.pyplot(fig)
#     if st.button("Анализировать диалог"):
#         st.write("Неоюбходимо выбрать файл для анализа")
        
