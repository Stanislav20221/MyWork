from docx import Document
import streamlit as st
from dotenv import load_dotenv
import openai
from openai import OpenAI
import os
import zipfile
import rarfile
import shutil
import docx
import re
import pandas as pd
import json
import matplotlib.pyplot as plt
import seaborn as sns
from io import BytesIO
load_dotenv()

client = OpenAI(
    api_key=os.getenv("OPENAI_API_KEY")

)
parazit= {"Неувернные формулировки (Я не знаю, Я не уверен, Насколько я помню, Если я не ошибаюсь, Так скажем , Скорее всего)", 
         "Отрицательные начала предложений (Нет, Не могу, Не знаю, Не получится), Фразы, которые могут вызвать сомнения: (К сожалению, Вы должны, Вам нужно, Это не моя компетенция)",
         "Сложные речевые конструкции: (Длинные и запутанные предложения, Использование сложных терминов без объяснения, Разговорные слова и выражения: (Смотрите, Типа, Ну, Эээ, Короче", "Уменьшительно-ласкательные формы: (Малыш, девочка, мальчик и т.д.)",
         "Междометия и слова-паразиты: (Эээ, ммм, угу, ага (в избытке)", "Ошибки в ударениях и произношении", "Неправильное произношение имен клиентов", "Грамматические ошибки", "Фразы, которые могут показаться высокомерными: (Я вас не понимаю, Вы неправы, Это не так",
         "Фразы, которые могут вызвать негативную реакцию: (Вы не правы, Это не моя проблема, Я не могу вам помочь)"}

parazittitle= {"Неуверенные формулировки", "Отрицательные начала предложений", "Сложные речевые конструкции", "Разговорные слова и выражения:", "Уменьшительно-ласкательные формы:",
         "Междометия и слова-паразиты", "Ошибки в ударениях и произношении", "Неправильное произношение имен клиентов", "Грамматические ошибки", "Фразы, которые могут показаться высокомерными",
         "Фразы, которые могут вызвать негативную реакцию"}


parazittitle1= {"Неувернные формулировки"}



system = '''Ты профессиональный нейроаналитик текста, умеешь находить слова-паразиты и нежелательные фразы. Твоя задача - найти указанные слова и фразы в тексте
        
'''
#==================================================================================== ФУНКЦИЯ ОТВЕТА CHAT GPT ======================================================================================================================================================
def answer_index(system, topic, dialog, temp, verbose=0):
    # Поиск релевантных отрезков из базы знаний по вопросу пользователя
    try:
        messages = [  # <-- Теперь тут 4 пробела
            {"role": "system", "content": system},
            {"role": "user", "content": f" Проанализируй текст диалога клиента и менеджера (Speaker 2 - это говорит клиент  Speaker 1 - это говорит менеджер) и выяви наличие нежелательных слов в диалоге менеждера. Выведи информацию о наличии/отсутствии и общего количества этих слов в диалоге менеджера по каждой категории. Результаты выведи в текстовом и JSON формате (для JSON формата выведи общее количество по каждой категории). Перед выводом каждого формата обозначь 'Текстовый формат:' 'JSON формат:' 'Список нежелательных слов: {topic} Текст: {dialog} Важно! Нежелательные слова и слова-паразиты выяви только в диалоге менеджера (Speaker 1). Тебе запрещено выявлять их в диалоге клиента. Ничего не придумывай от себя"}
        ]

        completion = client.chat.completions.create(  # <-- Уровень отступа теперь правильный
            model="gpt-4o-mini",
            messages=messages,
            temperature=temp
        )
        answer = completion.choices[0].message.content
        return answer
    except Exception as e:
        return "Ошибка доступа к Open AI. Проверьте API-ключ."


#==================================================================================== ФУНКЦИЯ ИЗВЛЕЧЕНИЯ JSON ======================================================================================================================================================    
def extract_json(text):
    match = re.search(r'\{[\s\S]*\}', text)  # Ищем JSON-блок в тексте
    if match:
        json_data = match.group(0)
        try:
            parsed_json = json.loads(json_data)
            return parsed_json  # Возвращаем как объект Python
        except json.JSONDecodeError:
            return "Ошибка: Некорректный JSON"
    return "Ошибка: JSON не найден"

def remove_json(text):
        # Удаляем всё после "JSON формат:" (включая сам JSON)
    text = re.sub(r'JSON формат:.*\{.*?\}', '', text, flags=re.DOTALL)

    return text.strip()  # Убираем лишние пробелы и переводы строк

#==================================================================================== ФУНКЦИЯ ОЧИСТКИ ТЕКСТА ОТ ЛИШНИХ СИМВОЛОВ ======================================================================================================================================================

def cleartext(text):
    report = re.sub(r"###.*?\n", "", text)
    text = re.sub(r"####.*?\n", "", text) 
    text = re.sub(r"\*\*", "", text)
    return report




st.title("Анализ текста на выявление запретных фраз и слов паразитов")
tab1, tab2,tab3 = st.tabs(["Диалог", "Результаты анализа (текст)", "Результаты анализа (таблица)"])
# Создание боковой панели
st.sidebar.header("Настройки модели")
temp = st.sidebar.slider("Температура", min_value=0.0, max_value=2.0, value=0.1) 
if "custom_css" not in st.session_state:
    st.session_state.custom_css = True  # Флаг, что CSS уже применён
    custom_text = "Перетащите файл сюда или выберите его вручную"
    st.markdown(f"""
        <style>
            .stFileUploader div div {{ visibility: hidden; }}
            .stFileUploader div div:after {{
                content: "{custom_text}";
                visibility: visible;
                display: block;
                text-align: center;
                font-size: 16px;
                color: #4F8BF9;
                font-weight: bold;
            }}
        </style>
    """, unsafe_allow_html=True)
 #                                         ЗАГРУЗКА ФАЙЛА
st.sidebar.header("Загрузить файл")
uploaded_file = st.sidebar.file_uploader("Выберите файл для загрузки", type=['docx'], key="file_uploader")
with tab1:
    if uploaded_file is not None:
        doc = Document(uploaded_file)
        content = "\n".join([para.text for para in doc.paragraphs])
        st.text_area("Содержимое файла:", content, height=300)           
with tab2:
    dicts = {}
    if uploaded_file is not None:
        doc = Document(uploaded_file)
        content = "\n".join([para.text for para in doc.paragraphs])
        dicts = answer_index(system, parazit, content, temp)
        st.markdown(remove_json(dicts))   
     
with tab3:
    if uploaded_file is not None:
        data = extract_json(dicts)      
        df=pd.DataFrame(list(data.items()), columns=["Тема", "Количество"])
        df.index.name = "№ п/п"
        st.table(df)